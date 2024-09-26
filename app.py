import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import Flask, request, render_template, send_file
import os
from concurrent.futures import ThreadPoolExecutor  # Import ThreadPoolExecutor for parallel processing

# Flask app setup
app = Flask(__name__)

# API key is now hardcoded
def configure_gemini_api():
    api_key = "AIzaSyDFlMKZPmXOtia0CbFZm5XyQWbAa1oy7YM"
    genai.configure(api_key=api_key)

def generate_text(prompt, max_tokens=None):
    model = genai.GenerativeModel('gemini-pro')
    generation_config = {}
    if max_tokens:
        generation_config['max_output_tokens'] = max_tokens

    response = model.generate_content(prompt, generation_config=generation_config)

    # Accessing the generated text from the response
    if response and response._result.candidates:
        generated_text = response._result.candidates[0].content.parts[0].text.strip()
        return generated_text.replace('*', '')  # Remove asterisks
    else:
        raise ValueError(f"Unexpected response structure: {response}")

def create_content(topic, structure=None, max_tokens=None):
    if not structure:
        outline_prompt = f"Generate a detailed outline for a document on: {topic}. Include an Introduction, 3-4 main sections, and a Conclusion."
        structure = generate_text(outline_prompt, max_tokens=max_tokens)

    sections = structure.split('\n')
    content = [("Understanding " + topic, "Introduction")]  # Main title and Introduction

    # Generate section content in parallel using ThreadPoolExecutor
    with ThreadPoolExecutor() as executor:
        future_to_section = {executor.submit(generate_text, f"Write detailed information about '{section}' related to {topic}. Include bullet points for key ideas."): section for section in sections if section.strip() and not section.startswith(('Introduction', 'Conclusion'))}

        for future in future_to_section:
            section = future_to_section[future]
            try:
                section_text = future.result()
                content.append((section, section_text))
            except Exception as exc:
                print(f"{section} generated an exception: {exc}")

    conclusion_prompt = f"Write a conclusion summarizing key points about {topic}."
    conclusion = generate_text(conclusion_prompt, max_tokens=max_tokens)
    content.append(("Conclusion", conclusion))

    return content

def create_word_document(content, filename='generated.docx'):
    doc = Document()

    # Configure document margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add main title
    title = doc.add_heading(content[0][0], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.size = Pt(24)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True

    for title, text in content[1:]:  # Skip the main title
        # Add section heading
        section_heading = doc.add_heading(title, level=1)
        section_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in section_heading.runs:
            run.font.size = Pt(24)
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.bold = True

        # Add content with formatting
        paragraphs = text.split('\n')
        for para in paragraphs:
            paragraph = doc.add_paragraph()
            apply_formatting(paragraph, para)
            for run in paragraph.runs:
                run.font.size = Pt(18)
                run.font.name = 'Arial'
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Save the document and handle exceptions
    try:
        doc.save(filename)
    except Exception as e:
        print(f"An error occurred while saving the document: {e}")

def apply_formatting(paragraph, text):
    clean_text = text.replace('*', '')
    if clean_text.strip().startswith('•'):
        paragraph.style = 'List Bullet'
        clean_text = clean_text.strip()[1:].strip()  # Remove the bullet point
    else:
        paragraph.style = 'Normal'
    paragraph.add_run(clean_text)

# Flask Routes
@app.route('/')
def index():
    return render_template('index.html')  # Load the HTML form

@app.route('/generate_document', methods=['POST'])
def generate_document():
    topic = request.form['topic']
    structure = request.form['structure']
    custom_structure = request.form['custom_structure'] if request.form['structure'].lower() == 'yes' else None
    num_pages = request.form['num_pages']

    max_tokens = int(num_pages) * 500 if num_pages else None

    # Create document content using your existing logic
    content = create_content(topic, structure=custom_structure, max_tokens=max_tokens)
    create_word_document(content, filename='generated.docx')

    # Return the generated document as a downloadable file
    return send_file('generated.docx', as_attachment=True)

# Main function for Flask
if __name__ == "__main__":
    configure_gemini_api()  # Initialize the API key
    app.run(debug=True)



